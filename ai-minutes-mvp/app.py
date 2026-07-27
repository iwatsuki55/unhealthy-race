import json
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import imageio_ffmpeg
import streamlit as st
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from openai import OpenAI, OpenAIError


APP_DIR = Path(__file__).parent
OUTPUT_DIR = APP_DIR / "outputs"
TERMS_PATH = APP_DIR / "terms_dictionary.json"
TEMPLATE_PATH = APP_DIR / "minutes_template.md"

DEFAULT_TITLE = "情報システム会議"
ALLOWED_EXTENSIONS = {".m4a", ".mp3", ".wav"}


load_dotenv(APP_DIR / ".env")


def get_env_int(name: str, default: int) -> int:
    value = os.getenv(name)
    if not value:
        return default
    try:
        return int(value)
    except ValueError:
        return default


TRANSCRIPTION_MODEL = os.getenv("OPENAI_TRANSCRIPTION_MODEL", "whisper-1")
MINUTES_MODEL = os.getenv("OPENAI_MINUTES_MODEL", "gpt-4o-mini")
CHUNK_LENGTH_MINUTES = get_env_int("CHUNK_LENGTH_MINUTES", 10)


def load_terms() -> Dict[str, str]:
    if not TERMS_PATH.exists():
        return {}

    try:
        with TERMS_PATH.open("r", encoding="utf-8") as file:
            data = json.load(file)
    except json.JSONDecodeError as exc:
        raise ValueError(f"terms_dictionary.json のJSON形式が正しくありません: {exc}") from exc

    if not isinstance(data, dict):
        raise ValueError("terms_dictionary.json は文字列キーと文字列値のオブジェクトにしてください。")

    return {str(key): str(value) for key, value in data.items()}


def apply_terms_dictionary(text: str, terms: Dict[str, str]) -> str:
    corrected = text
    for source, target in sorted(terms.items(), key=lambda item: len(item[0]), reverse=True):
        corrected = corrected.replace(source, target)
    return corrected


def detect_terms_in_text(text: str, terms: Dict[str, str]) -> List[str]:
    detected = set()
    for source, target in terms.items():
        if source in text or target in text:
            detected.add(target)
    return sorted(detected)


def save_uploaded_file(uploaded_file) -> Path:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(".m4a, .mp3, .wav の音声ファイルを選択してください。")

    temp_dir = Path(tempfile.mkdtemp(prefix="ai_minutes_"))
    file_path = temp_dir / uploaded_file.name
    with file_path.open("wb") as file:
        file.write(uploaded_file.getbuffer())
    return file_path


def split_audio(file_path: Path, chunk_length_minutes: int) -> List[Path]:
    chunk_length_seconds = max(1, chunk_length_minutes) * 60
    temp_dir = Path(tempfile.mkdtemp(prefix="ai_minutes_chunks_"))
    output_pattern = temp_dir / "chunk_%03d.wav"
    ffmpeg_path = imageio_ffmpeg.get_ffmpeg_exe()

    command = [
        ffmpeg_path,
        "-hide_banner",
        "-loglevel",
        "error",
        "-i",
        str(file_path),
        "-f",
        "segment",
        "-segment_time",
        str(chunk_length_seconds),
        "-reset_timestamps",
        "1",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-c:a",
        "pcm_s16le",
        str(output_pattern),
    ]

    result = subprocess.run(command, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        detail = result.stderr.strip() or result.stdout.strip() or "詳細不明"
        raise RuntimeError(f"音声ファイルを分割できませんでした。ファイル形式を確認してください。詳細: {detail}")

    chunk_paths = sorted(temp_dir.glob("chunk_*.wav"))
    if not chunk_paths:
        raise RuntimeError("音声ファイルの分割結果が空でした。音声ファイルが破損していないか確認してください。")

    return chunk_paths


def build_transcription_prompt(terms: Dict[str, str]) -> str:
    preferred_terms = sorted(set(terms.values()))
    terms_text = "、".join(preferred_terms)
    return (
        "これは病院・医療法人の情報システム会議の日本語音声です。"
        "発言内容を自然な日本語で正確に文字起こししてください。"
        f"専門用語は、実際に聞こえた場合のみ次の表記を優先してください: {terms_text}"
        "この一覧は表記補正用の候補であり、聞こえない用語を追加してはいけません。"
    )


def transcribe_chunks(client: OpenAI, chunk_paths: List[Path], terms: Dict[str, str]) -> str:
    prompt = build_transcription_prompt(terms)
    transcripts: List[str] = []

    for index, chunk_path in enumerate(chunk_paths, start=1):
        with chunk_path.open("rb") as audio_file:
            result = client.audio.transcriptions.create(
                model=TRANSCRIPTION_MODEL,
                file=audio_file,
                language="ja",
                prompt=prompt,
                response_format="text",
            )
        transcripts.append(f"【チャンク {index}】\n{str(result).strip()}")

    return "\n\n".join(transcripts)


def load_minutes_template(title: str, meeting_date: str) -> str:
    if TEMPLATE_PATH.exists():
        template = TEMPLATE_PATH.read_text(encoding="utf-8")
    else:
        template = "# {title}\n\n## 会議概要\n\n## 議題\n\n## 論点\n\n## 議論内容\n\n## 決定事項\n\n## 継続検討事項\n\n## アクションアイテム\n\n## 経営陣向けサマリー\n"
    return template.format(title=title, meeting_date=meeting_date)


def generate_minutes(client: OpenAI, transcript: str, title: str, meeting_date: str, terms: Dict[str, str]) -> str:
    template = load_minutes_template(title, meeting_date)
    detected_terms = detect_terms_in_text(transcript, terms)
    terms_text = "\n".join(f"- {term}" for term in detected_terms) if detected_terms else "- 該当なし"

    system_prompt = (
        "あなたは病院・医療法人の情報システム会議に詳しいコンサルタントです。"
        "文字起こしから、経営層が意思決定とリスクを把握しやすい議事録を作成してください。"
        "事実と推測を混同せず、音声にない担当者や期限は「未定」としてください。"
        "過度に口語的な表現を避け、簡潔で実務的な日本語にしてください。"
        "議題、論点、決定事項、継続検討事項は冗長にせず、経営層が一読で判断できる粒度に要約してください。"
        "文字起こしに明示されていない議題、製品名、決定事項、費用、期限を追加してはいけません。"
    )
    user_prompt = f"""
以下の文字起こしをもとに、指定テンプレートの構成を維持してMarkdown議事録を作成してください。

タイトル: {title}
会議日: {meeting_date}

文字起こし内で検出された優先表記の専門用語:
{terms_text}

作成方針:
- 見出し構成はテンプレートどおり維持する。
- 文字起こしに含まれる内容だけで作成する。テンプレート、専門用語辞書、過去の議事録から内容を補完しない。
- 「議題」は会議で扱ったテーマ名に絞る。
- 「論点」は判断・確認が必要な争点を短く列挙する。
- 「議論内容」はテーマ別の小見出しを置き、実際に話された内容を事実ベースで整理する。
- 「決定事項」は決まったことだけを書く。決まっていないことは入れない。
- 「継続検討事項」は次回以降に確認・検討すべき事項を書く。
- 「アクションアイテム」は担当者、期限、内容が読み取れる場合のみ表にする。読み取れない場合は「該当なし」とする。
- 「経営陣向けサマリー」はリスク、費用、運用影響、意思決定ポイントを3から5点で簡潔にまとめる。
- 話者名や人名は、発言者ではなく検討対象や担当者として意味がある場合のみ使う。
- 聞き取りが不確かな固有名詞は、上記の「検出された優先表記」に含まれる場合のみその表記を優先する。

テンプレート:
{template}

文字起こし:
{transcript}
"""

    response = client.chat.completions.create(
        model=MINUTES_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    content = response.choices[0].message.content
    if not content:
        raise RuntimeError("議事録生成の応答が空でした。")
    return content.strip()


def markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Hiragino Sans"
    styles["Normal"].font.size = Pt(10.5)

    in_table = False
    table_rows: List[List[str]] = []

    def flush_table() -> None:
        nonlocal table_rows, in_table
        if len(table_rows) >= 2:
            header = table_rows[0]
            body = table_rows[2:] if is_markdown_separator(table_rows[1]) else table_rows[1:]
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for col_index, value in enumerate(header):
                table.rows[0].cells[col_index].text = value
            for row in body:
                cells = table.add_row().cells
                for col_index, value in enumerate(row[: len(header)]):
                    cells[col_index].text = value
        table_rows = []
        in_table = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()

        if line.startswith("|") and line.endswith("|"):
            in_table = True
            table_rows.append(parse_markdown_table_row(line))
            continue

        if in_table:
            flush_table()

        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)

    if in_table:
        flush_table()

    document.save(output_path)


def parse_markdown_table_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip("|").split("|")]


def is_markdown_separator(row: List[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) is not None for cell in row)


def make_safe_filename_part(value: str, fallback: str) -> str:
    safe_value = re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" ._")
    return safe_value or fallback


def save_outputs(
    minutes_markdown: str,
    corrected_transcript: str,
    title: str,
    meeting_date: str,
    source_filename: str,
) -> Dict[str, Path]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    date_for_filename = meeting_date.replace("-", "")
    run_stamp = datetime.now().strftime("%H%M%S")
    safe_title = make_safe_filename_part(title, DEFAULT_TITLE)
    source_stem = make_safe_filename_part(Path(source_filename).stem, "音声")
    base_name = f"{safe_title}_{date_for_filename}_{run_stamp}_{source_stem}"
    run_output_dir = OUTPUT_DIR / f"{date_for_filename}_{run_stamp}_{source_stem}"
    run_output_dir.mkdir(exist_ok=True)
    md_path = run_output_dir / f"{base_name}.md"
    docx_path = run_output_dir / f"{base_name}.docx"
    transcript_path = run_output_dir / f"{base_name}_transcript.txt"

    md_path.write_text(minutes_markdown, encoding="utf-8")
    transcript_path.write_text(corrected_transcript, encoding="utf-8")
    markdown_to_docx(minutes_markdown, docx_path)

    return {"markdown": md_path, "docx": docx_path, "transcript": transcript_path}


def render_openai_error(exc: OpenAIError) -> None:
    error_text = str(exc)
    status_code = getattr(exc, "status_code", None)
    error_code = getattr(exc, "code", None)

    if status_code == 429 or "insufficient_quota" in error_text:
        st.error(
            "OpenAI APIの利用枠または課金設定が不足しています。"
            "OpenAI PlatformのBillingで支払い方法、利用上限、残高を確認してください。"
        )
        st.markdown(
            "- [Billing overview](https://platform.openai.com/settings/organization/billing/overview)\n"
            "- [Usage](https://platform.openai.com/usage)\n"
            "- [API keys](https://platform.openai.com/api-keys)"
        )
        st.caption("ChatGPT Plusの契約とは別に、OpenAI API側の課金設定が必要な場合があります。")
        return

    if status_code == 401 or "invalid_api_key" in error_text:
        st.error(".env の OPENAI_API_KEY が正しいか確認してください。")
        st.markdown("[API keys](https://platform.openai.com/api-keys)")
        return

    if status_code == 429 or error_code == "rate_limit_exceeded":
        st.error("OpenAI APIのレート制限に達しました。少し時間を置いてから再実行してください。")
        return

    st.error(f"OpenAI APIでエラーが発生しました: {exc}")


def render_sidebar() -> None:
    st.sidebar.header("設定")
    st.sidebar.write(f"文字起こしモデル: `{TRANSCRIPTION_MODEL}`")
    st.sidebar.write(f"議事録生成モデル: `{MINUTES_MODEL}`")
    st.sidebar.write(f"分割単位: `{CHUNK_LENGTH_MINUTES}` 分")
    st.sidebar.write("保存方式: `実行ごとの個別フォルダ`")
    st.sidebar.caption("モデルや分割単位は .env で変更できます。")


def main() -> None:
    st.set_page_config(page_title="AI議事録作成ツール MVP", layout="wide")
    render_sidebar()

    st.title("AI議事録作成ツール MVP")
    st.caption("病院・医療法人の情報システム会議向けに、音声から経営層向け議事録を生成します。")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error(".env に OPENAI_API_KEY を設定してください。")
        st.stop()

    with st.form("minutes_form"):
        title = st.text_input("議事録タイトル", value=DEFAULT_TITLE)
        meeting_date = st.date_input("会議日", value=datetime.now().date())
        uploaded_file = st.file_uploader(
            "音声ファイルを選択またはドラッグ＆ドロップ",
            type=["m4a", "mp3", "wav"],
            accept_multiple_files=False,
        )
        submitted = st.form_submit_button("議事録を作成")

    st.caption("出力は毎回 `outputs/YYYYMMDD_HHMMSS_元音声名/` に保存され、過去の議事録を上書きしません。")

    if not submitted:
        st.info(".m4a, .mp3, .wav の会議音声をアップロードして開始してください。")
        return

    if uploaded_file is None:
        st.warning("音声ファイルを選択してください。")
        return

    client = OpenAI(api_key=api_key)
    meeting_date_text = meeting_date.strftime("%Y-%m-%d")

    try:
        with st.status("処理を開始しています...", expanded=True) as status:
            st.write("専門用語辞書を読み込み中...")
            terms = load_terms()

            st.write("アップロードファイルを一時保存中...")
            audio_path = save_uploaded_file(uploaded_file)

            st.write("音声を分割中...")
            chunk_paths = split_audio(audio_path, CHUNK_LENGTH_MINUTES)
            st.write(f"{len(chunk_paths)} 個のチャンクに分割しました。")

            st.write("OpenAI APIで文字起こし中...")
            raw_transcript = transcribe_chunks(client, chunk_paths, terms)

            st.write("専門用語辞書で表記を補正中...")
            corrected_transcript = apply_terms_dictionary(raw_transcript, terms)

            st.write("経営層向け議事録を生成中...")
            minutes_markdown = generate_minutes(
                client=client,
                transcript=corrected_transcript,
                title=title,
                meeting_date=meeting_date_text,
                terms=terms,
            )
            minutes_markdown = apply_terms_dictionary(minutes_markdown, terms)

            st.write("Markdownとdocxを保存中...")
            output_paths = save_outputs(
                minutes_markdown=minutes_markdown,
                corrected_transcript=corrected_transcript,
                title=title,
                meeting_date=meeting_date_text,
                source_filename=uploaded_file.name,
            )

            status.update(label="議事録の作成が完了しました。", state="complete")

        st.subheader("生成された議事録")
        st.markdown(minutes_markdown)

        with st.expander("文字起こし結果を確認する"):
            st.text_area("補正後の文字起こし", corrected_transcript, height=300)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "Markdownをダウンロード",
                data=output_paths["markdown"].read_bytes(),
                file_name=output_paths["markdown"].name,
                mime="text/markdown",
            )
        with col2:
            st.download_button(
                "Word（docx）をダウンロード",
                data=output_paths["docx"].read_bytes(),
                file_name=output_paths["docx"].name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        st.success(f"今回の保存先: {output_paths['markdown'].parent}")
        st.caption(f"文字起こし保存先: {output_paths['transcript']}")

    except OpenAIError as exc:
        render_openai_error(exc)
    except FileNotFoundError as exc:
        st.error(f"必要なファイルが見つかりません: {exc}")
    except ValueError as exc:
        st.error(str(exc))
    except RuntimeError as exc:
        st.error(str(exc))
    except Exception as exc:
        st.error(f"予期しないエラーが発生しました: {type(exc).__name__}: {exc}")


if __name__ == "__main__":
    main()
